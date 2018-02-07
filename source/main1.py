# -*- coding: utf-8 -*-
"""
Created on Fri Jan 12 09:55:52 2018

@author: user
"""

from docx import Document
import json
import table
document = Document() 
def fun(document):
    paras = document.paragraphs
    lcount = []  # 存储所有段落的对象 
  
    for i in range(len(paras)):
        if paras[i].style.name == 'Heading 2':
            lcount.append(i)
#            print(paras[i].text)
#    print(len(lcount))      
    # 提取每一段的数据
    for i in range(len(lcount)):
        if i+1 >= len(lcount):
            break
#       print(lcount)

        ax = lcount[i]
        bx = lcount[i+1]
#       print(ax,bx)
        Head4_list = []
        Head3_list = []
        normal_list = []
        # 取出当前一级标题需要的n个段落
        for j in range(int(ax),int(bx),1):
#            print(paras[j].text)
            # 获取并打印二级标题
            if paras[j].style.name == 'Heading 2':
                Head2 = paras[j].text
#                print(Head2)
            # 获取并打印三级题
            if paras[j].style.name == 'Heading 3':
                Head3 =  paras[j].text
#                print(Head3)
                Head3_list.append(Head3)
                
              # 获取并打印四级题
            elif paras[j].style.name == 'Heading 4':
                Head4 =  paras[j].text
#                print(Head4)
                Head4_list.append(Head4)

            # 获取并打印正文
            elif paras[j].style.name == 'Normal':
                normal = paras[j].text                
                new_normal = normal.replace('\u3000', ' ')
                normal_list.append(new_normal)               
#                   
        break


    data = {Head2:{Head3_list[0]:{Head4_list[0]:normal_list[0],
                                       Head4_list[1]:normal_list[1]+normal_list[2]+normal_list[3],
                                       Head4_list[2]:normal_list[4]+normal_list[5],
                                       Head4_list[3]:'',
                                       Head4_list[4]:table.table(0,document)}
       
            }}
    print(data)
   
    with open('gf_asset_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
        
        
    data = {Head2:{Head3_list[1]:{Head4_list[5]:normal_list[6],
                                       Head4_list[6]:normal_list[7]+normal_list[8]+normal_list[9],
                                       Head4_list[7]:normal_list[10]+normal_list[11],
                                       Head4_list[8]:'',
                                       Head4_list[9]:table.table(1,document)}
       
            }}
    print(data)
   
    with open('gf_prod_asset_rel.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
   
   
fun(Document('理财监管新规数据库优化11.9.docx')) 
    
    
    