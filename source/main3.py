# -*- coding: utf-8 -*-
"""
Created on Fri Jan 12 09:55:52 2018

@author: user
"""

from docx import Document
import json
import table
document = Document() 
def fun(document):
    paras = document.paragraphs
    lcount = []  # 存储所有段落的对象 
  
    for i in range(len(paras)):
        if paras[i].style.name == 'Heading 2':
            lcount.append(i)     
    # 提取每一段的数据
    for i in range(len(lcount)):
        if i+1 >= len(lcount):
            break
        ax = lcount[i+2]
        bx = lcount[i+3]
#       print(ax,bx)
        Head4_list = []
        Head3_list = []
        normal_list = []
        # 取出当前一级标题需要的n个段落
        for j in range(int(ax),int(bx),1):
#            print(paras[j].text)
            # 获取并打印二级标题
            if paras[j].style.name == 'Heading 2':
                Head2 = paras[j].text
                print(Head2)
            # 获取并打印三级题
            if paras[j].style.name == 'Heading 3':
                Head3 =  paras[j].text
                print(Head3)
                Head3_list.append(Head3)               
              # 获取并打印四级题
            elif paras[j].style.name == 'Heading 4':
                Head4 =  paras[j].text
                print(Head4)
                Head4_list.append(Head4)
            # 获取并打印正文
            elif paras[j].style.name == 'Normal':
                normal = paras[j].text                
                print(normal)
                new_normal = normal.replace('\u3000', ' ')
                normal_list.append(new_normal)                                    
        break
#    print(Head3_list)
    data = {Head2:{Head3_list[0]:{Head4_list[0]:normal_list[0],
                                       Head4_list[1]:normal_list[1]+normal_list[2]+normal_list[3],
                                       Head4_list[2]:normal_list[4]+normal_list[5],
                                       Head4_list[3]:'',
                                       Head4_list[4]:table.table(5,document)}       
            }}
    print(data)
    with open('gf_prod_register_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[1]:{Head4_list[5]:normal_list[6],
                                       Head4_list[6]:normal_list[7]+normal_list[8]+normal_list[9],
                                       Head4_list[7]:normal_list[10]+normal_list[11],
                                       Head4_list[8]:'',
                                       Head4_list[9]:table.table(6,document)}
            }}
    print(data)  
    with open('gf_prod_brand_info','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)        
    data = {Head2:{Head3_list[2]:{Head4_list[10]:normal_list[12],
                                       Head4_list[11]:normal_list[13]+normal_list[14]+normal_list[15],
                                       Head4_list[12]:normal_list[16]+normal_list[17],
                                       Head4_list[13]:'',
                                       Head4_list[14]:table.table(7,document)}       
            }}
    print(data)  
    with open('gf_prod_issue_area.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[3]:{Head4_list[15]:normal_list[18],
                                       Head4_list[16]:normal_list[19]+normal_list[20]+normal_list[21],
                                       Head4_list[17]:normal_list[22]+normal_list[23],
                                       Head4_list[18]:'',
                                       Head4_list[19]:table.table(8,document)}       
            }}
    print(data)  
    with open('ta_retailer_limit_ctl.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[4]:{Head4_list[20]:normal_list[24],
                                       Head4_list[21]:normal_list[25]+normal_list[26]+normal_list[27],
                                       Head4_list[22]:normal_list[28]+normal_list[29],
                                       Head4_list[23]:'',
                                       Head4_list[24]:table.table(9,document)}       
            }}
    print(data)  
    with open('gf_limit_ctl.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[5]:{Head4_list[25]:normal_list[30],
                                       Head4_list[26]:normal_list[31]+normal_list[32]+normal_list[33],
                                       Head4_list[27]:normal_list[34]+normal_list[35],
                                       Head4_list[28]:'',
                                       Head4_list[29]:table.table(10,document)}       
            }}
    print(data)  
    with open('gf_prod_base_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[6]:{Head4_list[30]:normal_list[36],
                                       Head4_list[31]:normal_list[37]+normal_list[38]+normal_list[39],
                                       Head4_list[32]:normal_list[40]+normal_list[41],
                                       Head4_list[33]:'',
                                       Head4_list[34]:table.table(11,document)}       
            }}
    print(data)  
    with open('gf_prod_organ_white_list.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[7]:{Head4_list[35]:normal_list[42],
                                       Head4_list[36]:normal_list[43]+normal_list[44]+normal_list[45],
                                       Head4_list[37]:normal_list[46]+normal_list[47],
                                       Head4_list[38]:'',
                                       Head4_list[39]:table.table(12,document)}       
            }}
    print(data)  
    with open('gf_prod_cust_white_list.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[8]:{Head4_list[40]:normal_list[48],
                                       Head4_list[41]:normal_list[49]+normal_list[50]+normal_list[51],
                                       Head4_list[42]:normal_list[52]+normal_list[53],
                                       Head4_list[43]:'',
                                       Head4_list[44]:table.table(13,document)}       
            }}
    print(data)  
    with open('gf_prod_point_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[9]:{Head4_list[45]:normal_list[54],
                                       Head4_list[46]:normal_list[55]+normal_list[56]+normal_list[57],
                                       Head4_list[47]:normal_list[58]+normal_list[59],
                                       Head4_list[48]:'',
                                       Head4_list[49]:table.table(14,document)}       
            }}
    print(data)  
    with open('gf_prod_pos_parameters.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[10]:{Head4_list[50]:normal_list[60],
                                       Head4_list[51]:normal_list[61]+normal_list[62]+normal_list[63],
                                       Head4_list[52]:normal_list[64]+normal_list[65],
                                       Head4_list[53]:'',
                                       Head4_list[54]:table.table(15,document)}       
            }}
    print(data)  
    with open('gf_prod_cust_prop.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[11]:{Head4_list[55]:normal_list[66],
                                       Head4_list[56]:normal_list[67]+normal_list[68]+normal_list[69],
                                       Head4_list[57]:normal_list[70]+normal_list[71],
                                       Head4_list[58]:'',
                                       Head4_list[59]:table.table(16,document)}       
            }}
    print(data)  
    with open('gf_reservation_buy_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[12]:{Head4_list[60]:normal_list[72],
                                       Head4_list[61]:normal_list[73]+normal_list[74]+normal_list[75],
                                       Head4_list[62]:normal_list[76]+normal_list[77],
                                       Head4_list[63]:'',
                                       Head4_list[64]:table.table(17,document)}       
            }}
    print(data)  
    with open('gf_prod_buy_infol.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[13]:{Head4_list[65]:normal_list[78],
                                       Head4_list[66]:normal_list[79]+normal_list[80]+normal_list[81],
                                       Head4_list[67]:normal_list[82]+normal_list[83],
                                       Head4_list[68]:'',
                                       Head4_list[69]:table.table(18,document)}       
            }}
    print(data)  
    with open('gf_prod_open_duration_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[14]:{Head4_list[70]:normal_list[84],
                                       Head4_list[71]:normal_list[85]+normal_list[86]+normal_list[87],
                                       Head4_list[72]:normal_list[88]+normal_list[89],
                                       Head4_list[73]:'',
                                       Head4_list[74]:table.table(19,document)}       
            }}
    print(data)  
    with open('gf_prod_redeem_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[15]:{Head4_list[75]:normal_list[90],
                                       Head4_list[76]:normal_list[91]+normal_list[92]+normal_list[93],
                                       Head4_list[77]:normal_list[94]+normal_list[95],
                                       Head4_list[78]:'',
                                       Head4_list[79]:table.table(20,document)}       
            }}
    print(data)  
    with open('gf_prod_dividend_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[16]:{Head4_list[80]:normal_list[96],
                                       Head4_list[81]:normal_list[97]+normal_list[98]+normal_list[99],
                                       Head4_list[82]:normal_list[100]+normal_list[101],
                                       Head4_list[83]:'',
                                       Head4_list[84]:table.table(21,document)}       
            }}
    print(data)  
    with open('gf_prod_netvalue_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[17]:{Head4_list[85]:normal_list[102],
                                       Head4_list[86]:normal_list[103]+normal_list[104]+normal_list[105],
                                       Head4_list[87]:normal_list[106]+normal_list[107],
                                       Head4_list[88]:'',
                                       Head4_list[89]:table.table(22,document)}       
            }}
    print(data)  
    with open('gf_prod_fix_profit_rate.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[18]:{Head4_list[90]:normal_list[108],
                                       Head4_list[91]:normal_list[109]+normal_list[110]+normal_list[111],
                                       Head4_list[92]:normal_list[112]+normal_list[113],
                                       Head4_list[93]:'',
                                       Head4_list[94]:table.table(23,document)}       
            }}
    print(data)  
    with open('gf_prod_fix_period_profit_rate.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[19]:{Head4_list[95]:normal_list[114],
                                       Head4_list[96]:normal_list[115]+normal_list[116]+normal_list[117],
                                       Head4_list[97]:normal_list[118]+normal_list[119],
                                       Head4_list[98]:'',
                                       Head4_list[99]:table.table(24,document)}       
            }}
    print(data)  
    with open('gf_prod_scale_profit_rate.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[20]:{Head4_list[100]:normal_list[120],
                                       Head4_list[101]:normal_list[121]+normal_list[122]+normal_list[123],
                                       Head4_list[102]:normal_list[124]+normal_list[125],
                                       Head4_list[103]:'',
                                       Head4_list[104]:table.table(25,document)}       
            }}
    print(data)  
    with open('gf_prod_fel_period_profit_rate.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[21]:{Head4_list[105]:normal_list[126],
                                       Head4_list[106]:normal_list[127]+normal_list[128]+normal_list[129],
                                       Head4_list[107]:normal_list[130]+normal_list[131],
                                       Head4_list[108]:'',
                                       Head4_list[109]:table.table(26,document)}       
            }}
    print(data)  
    with open('gf_prod_fix_sell_fee_rate.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[22]:{Head4_list[110]:normal_list[132],
                                       Head4_list[111]:normal_list[133]+normal_list[134]+normal_list[135],
                                       Head4_list[112]:normal_list[136]+normal_list[137],
                                       Head4_list[113]:'',
                                       Head4_list[114]:table.table(27,document)}       
            }}
    print(data)  
    with open('gf_prod_fel_sell_fee_rate.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[23]:{Head4_list[115]:normal_list[138],
                                       Head4_list[116]:normal_list[139]+normal_list[140]+normal_list[141],
                                       Head4_list[117]:normal_list[142]+normal_list[143],
                                       Head4_list[118]:'',
                                       Head4_list[119]:table.table(28,document)}       
            }}
    print(data)  
    with open('gf_prod_statistics_para.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[24]:{Head4_list[120]:normal_list[144],
                                       Head4_list[121]:normal_list[145]+normal_list[146]+normal_list[147],
                                       Head4_list[122]:normal_list[148]+normal_list[149],
                                       Head4_list[123]:'',
                                       Head4_list[124]:table.table(29,document)}       
            }}
    print(data)  
    with open('gf_prod_profit_detail.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[25]:{Head4_list[125]:normal_list[150],
                                       Head4_list[126]:normal_list[151]+normal_list[152]+normal_list[153],
                                       Head4_list[127]:normal_list[154]+normal_list[155],
                                       Head4_list[128]:'',
                                       Head4_list[129]:table.table(30,document)}       
            }}
    print(data)  
    with open('gf_prod_redeem_detail.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[26]:{Head4_list[130]:normal_list[156],
                                       Head4_list[131]:normal_list[157]+normal_list[158]+normal_list[159],
                                       Head4_list[132]:normal_list[160]+normal_list[161],
                                       Head4_list[133]:'',
                                       Head4_list[134]:table.table(31,document)}       
            }}
    print(data)  
    with open('gf_prod_LIMIT_detail.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[27]:{Head4_list[135]:normal_list[162],
                                       Head4_list[136]:normal_list[163]+normal_list[164]+normal_list[165],
                                       Head4_list[137]:normal_list[166]+normal_list[167],
                                       Head4_list[138]:'',
                                       Head4_list[139]:table.table(32,document)}       
            }}
    print(data)  
    with open('gf_prod_open_duration_rule.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[28]:{Head4_list[140]:normal_list[168],
                                       Head4_list[141]:normal_list[169]+normal_list[170]+normal_list[171],
                                       Head4_list[142]:normal_list[172]+normal_list[173],
                                       Head4_list[143]:'',
                                       Head4_list[144]:table.table(33,document)}       
            }}
    print(data)  
    with open('gf_prod_profit_detail_sum.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)

   
   
fun(Document('理财监管新规数据库优化11.9.docx')) 
    
    
    


