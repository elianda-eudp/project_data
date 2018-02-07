# -*- coding: utf-8 -*-
"""
Created on Fri Jan 12 09:55:52 2018

@author: user
"""

from docx import Document
import json
import table
document = Document() 
def fun(document):
    paras = document.paragraphs
    lcount = []  # 存储所有段落的对象 
  
    for i in range(len(paras)):
        if paras[i].style.name == 'Heading 2':
            lcount.append(i)     
    # 提取每一段的数据
    for i in range(len(lcount)):
        if i+1 >= len(lcount):
            break
        ax = lcount[i+6]
        bx = lcount[i+7]
#       print(ax,bx)
        Head4_list = []
        Head3_list = []
        normal_list = []
        # 取出当前一级标题需要的n个段落
        for j in range(int(ax),int(bx),1):
#            print(paras[j].text)
            # 获取并打印二级标题
            if paras[j].style.name == 'Heading 2':
                Head2 = paras[j].text
                print(Head2)
            # 获取并打印三级题
            if paras[j].style.name == 'Heading 3':
                Head3 =  paras[j].text
                print(Head3)
                Head3_list.append(Head3)               
              # 获取并打印四级题
            elif paras[j].style.name == 'Heading 4':
                Head4 =  paras[j].text
                print(Head4)
                Head4_list.append(Head4)
            # 获取并打印正文
            elif paras[j].style.name == 'Normal':
                normal = paras[j].text                
                print(normal)
                new_normal = normal.replace('\u3000', ' ')
                normal_list.append(new_normal)                                  
        break

    data = {Head2:{Head3_list[0]:{Head4_list[0]:normal_list[0],
                                       Head4_list[1]:normal_list[1]+normal_list[2]+normal_list[3],
                                       Head4_list[2]:normal_list[4]+normal_list[5],
                                       Head4_list[3]:'',
                                       Head4_list[4]:table.table(40,document)}       
            }}
    print(data)
    with open('gf_prod_auto_ctl_time_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[1]:{Head4_list[5]:normal_list[6],
                                       Head4_list[6]:normal_list[7]+normal_list[8]+normal_list[9],
                                       Head4_list[7]:normal_list[10]+normal_list[11],
                                       Head4_list[8]:'',
                                       Head4_list[9]:table.table(41,document)}
            }}
    print(data)  
    with open('gf_prod_auto_open_info','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)        
    data = {Head2:{Head3_list[2]:{Head4_list[10]:normal_list[12],
                                       Head4_list[11]:normal_list[13]+normal_list[14]+normal_list[15],
                                       Head4_list[12]:normal_list[16]+normal_list[17],
                                       Head4_list[13]:'',
                                       Head4_list[14]:table.table(42,document)}       
            }}
    print(data)  
    with open('gf_prod_timing_buy_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[3]:{Head4_list[15]:normal_list[18],
                                       Head4_list[16]:normal_list[19]+normal_list[20]+normal_list[21],
                                       Head4_list[17]:normal_list[22]+normal_list[23],
                                       Head4_list[18]:'',
                                       Head4_list[19]:table.table(43,document)}       
            }}
    print(data)  
    with open('gf_prod_order_auto_buy_agree.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[4]:{Head4_list[20]:normal_list[24],
                                       Head4_list[21]:normal_list[25]+normal_list[26]+normal_list[27],
                                       Head4_list[22]:normal_list[28]+normal_list[29],
                                       Head4_list[23]:'',
                                       Head4_list[24]:table.table(44,document)}       
            }}
    print(data)  
    with open('gf_prod_order_redeem_agree.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[5]:{Head4_list[25]:normal_list[30],
                                       Head4_list[26]:normal_list[31]+normal_list[32]+normal_list[33],
                                       Head4_list[27]:normal_list[34]+normal_list[35],
                                       Head4_list[28]:table.table(45,document),
                                       Head4_list[29]:table.table(46,document)}       
            }}
    print(data)  
    with open('gf_prod_order_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[6]:{Head4_list[30]:normal_list[36],
                                       Head4_list[31]:normal_list[37]+normal_list[38]+normal_list[39],
                                       Head4_list[32]:normal_list[40]+normal_list[41],
                                       Head4_list[33]:'',
                                       Head4_list[34]:table.table(47,document)}       
            }}
    print(data)  
    with open('gf_cust_unconfirm_buy_amt.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[7]:{Head4_list[35]:normal_list[42],
                                       Head4_list[36]:normal_list[43]+normal_list[44]+normal_list[45],
                                       Head4_list[37]:normal_list[46]+normal_list[47],
                                       Head4_list[38]:'',
                                       Head4_list[39]:table.table(48,document)}       
            }}
    print(data)  
    with open('gf_huge_redeem_vastly_ctl.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[8]:{Head4_list[40]:normal_list[48],
                                       Head4_list[41]:normal_list[49]+normal_list[50]+normal_list[51],
                                       Head4_list[42]:normal_list[52]+normal_list[53],
                                       Head4_list[43]:'',
                                       Head4_list[44]:table.table(49,document)}       
            }}
    print(data)  
    with open('gf_prod_pos_accin_list.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[9]:{Head4_list[45]:normal_list[54],
                                       Head4_list[46]:normal_list[55]+normal_list[56]+normal_list[57],
                                       Head4_list[47]:normal_list[58]+normal_list[59],
                                       Head4_list[48]:'',
                                       Head4_list[49]:table.table(50,document)}       
            }}
    print(data)  
    with open('gf_prod_pos_balance.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[10]:{Head4_list[50]:normal_list[60],
                                       Head4_list[51]:normal_list[61]+normal_list[62]+normal_list[63],
                                       Head4_list[52]:normal_list[64]+normal_list[65],
                                       Head4_list[53]:'',
                                       Head4_list[54]:table.table(51,document)}       
            }}
    print(data)  
    with open('gf_prod_quot_pledge_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[11]:{Head4_list[55]:normal_list[66],
                                       Head4_list[56]:normal_list[67]+normal_list[68]+normal_list[69],
                                       Head4_list[57]:normal_list[70]+normal_list[71],
                                       Head4_list[58]:'',
                                       Head4_list[59]:table.table(52,document)}       
            }}
    print(data)  
    with open('gf_prod_asset_prove_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[12]:{Head4_list[60]:normal_list[72],
                                       Head4_list[61]:normal_list[73]+normal_list[74]+normal_list[75],
                                       Head4_list[62]:normal_list[76]+normal_list[77],
                                       Head4_list[63]:'',
                                       Head4_list[64]:table.table(53,document)}       
            }}
    print(data)  
    with open('gf_endday_trade_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[13]:{Head4_list[65]:normal_list[78],
                                       Head4_list[66]:normal_list[79]+normal_list[80]+normal_list[81],
                                       Head4_list[67]:normal_list[82]+normal_list[83],
                                       Head4_list[68]:table.table(54,document),
                                       Head4_list[69]:table.table(55,document)}       
            }}
    print(data)  
    with open('gf_prod_freeze_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[14]:{Head4_list[70]:normal_list[84],
                                       Head4_list[71]:normal_list[85]+normal_list[86]+normal_list[87],
                                       Head4_list[72]:normal_list[88]+normal_list[89],
                                       Head4_list[73]:'',
                                       Head4_list[74]:table.table(56,document)}       
            }}
    print(data)  
    with open('gf_no_trade_transfer_info.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[15]:{Head4_list[75]:normal_list[90],
                                       Head4_list[76]:normal_list[91]+normal_list[92]+normal_list[93],
                                       Head4_list[77]:normal_list[94]+normal_list[95],
                                       Head4_list[78]:table.table(57,document),
                                       Head4_list[79]:table.table(58,document)}       
            }}
    print(data)  
    with open('gf_sell_trade_list_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[16]:{Head4_list[80]:normal_list[96],
                                       Head4_list[81]:normal_list[97]+normal_list[98]+normal_list[99],
                                       Head4_list[82]:normal_list[100]+normal_list[101],
                                       Head4_list[83]:'',
                                       Head4_list[84]:table.table(59,document)}       
            }}
    print(data)  
    with open('gf_prod_cancell_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[17]:{Head4_list[85]:normal_list[102],
                                       Head4_list[86]:normal_list[103]+normal_list[104]+normal_list[105],
                                       Head4_list[87]:normal_list[106]+normal_list[107],
                                       Head4_list[88]:'',
                                       Head4_list[89]:table.table(60,document)}       
            }}
    print(data)  
    with open('gf_checklist_point.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
    data = {Head2:{Head3_list[18]:{Head4_list[90]:normal_list[108],
                                       Head4_list[91]:normal_list[109]+normal_list[110]+normal_list[111],
                                       Head4_list[92]:normal_list[112]+normal_list[113],
                                       Head4_list[93]:table.table(61,document),
                                       Head4_list[94]:table.table(62,document)}       
            }}
    print(data)  
    with open('gf_prod_transf_note.json','w+') as f:
        str_data = json.dumps(data)
        f.write(str_data)
   
   
fun(Document('理财监管新规数据库优化11.9.docx')) 
    
    
    








