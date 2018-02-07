from docx import Document
document = Document() 
def table(num,document):

    row1_list = []
    row_content = []
    """获取第一个表格内容"""
    table = document.tables[num]
    """获取第一行标题"""
    row_1 = table.row_cells(0)
    """将第一行标题提取出来"""
    for row in table.rows:
        for cell1 in row.cells:
    #            print(cell.text)
            for cell2 in row_1:
#                print(cell2.text)
                if cell1.text == cell2.text:
                    key = cell1.text                 
                    row1_list.append(key)
#                    print(key)
                    break
            else:
                """将内容提取出来"""
                value = cell1.text
                new_value = value.replace('\n', ' ')
                row_content.append(new_value)

    dict1={}

    for i in range(len(row1_list)): # 循环标题列表
        dict2={}
        a_list=[j for j in range(i,len(row_content),len(row1_list)) ] # 循环内容列表
        for k in range(len(a_list)):  
            dict2[row1_list[i]+str(k)] = row_content[a_list[k]] # 构造内层字典
        dict1[row1_list[i]] = dict2     # 构造整体字典
              
#    print(dict1)
    return dict1
    
     
   
                
table(0,Document('理财监管新规数据库优化11.9.docx'))
#table(1)